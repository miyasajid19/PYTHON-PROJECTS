import os
import shutil
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

def clean_text(text):
    return re.sub(r'[^\x09\x0A\x0D\x20-\x7F\xA0-\uD7FF\uE000-\uFFFD]', '', text)

def add_table_of_contents(doc):
    doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph('TOC will be generated here after you update the fields in Word.')
    doc.add_page_break()

def add_page_numbers(doc):
    sections = doc.sections
    
    for i, section in enumerate(sections):
        footer = section.footer
        if i == 0:
            continue
        paragraph = footer.add_paragraph()
        paragraph.add_run('Page ').add_run().add_break()

def set_document_properties(doc, author_name):
    doc.core_properties.author = author_name

def create_document_from_folder(folder_path, cover_path):
    folder_name = os.path.basename(folder_path)
    
    cpp_files = [file for file in os.listdir(folder_path) if file.endswith('.cpp')]
    if not cpp_files:
        print(f'No .cpp files found in {folder_path}. Skipping.')
        return
    
    temp_doc_path = os.path.join(folder_path, f'{folder_name}_Temp.docx')
    try:
        shutil.copy(cover_path, temp_doc_path)
    except Exception as e:
        print(f'Error copying cover page file: {e}')
        return

    try:
        doc = Document(temp_doc_path)
    except Exception as e:
        print(f'Error opening temporary document: {e}')
        return

    add_table_of_contents(doc)

    for file in cpp_files:
        file_path = os.path.join(folder_path, file)
        if os.path.isfile(file_path):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        content = f.read()
                except Exception as e:
                    print(f'Error reading file {file}: {e}')
                    continue

            cleaned_content = clean_text(content)
            doc.add_heading(f'Solution from {file}', level=1)
            doc.add_paragraph(cleaned_content)

    add_page_numbers(doc)
    set_document_properties(doc, "Sajid Miya")

    output_path = os.path.join(folder_path, f'{folder_name}_Solutions.docx')
    try:
        doc.save(output_path)
        print(f'Document saved as {output_path}')
    except Exception as e:
        print(f'Error saving document: {e}')

    if os.path.exists(temp_doc_path):
        os.remove(temp_doc_path)
        print(f'Temporary file {temp_doc_path} deleted.')

def process_all_folders(base_path, cover_path):
    for root, dirs, files in os.walk(base_path):
        for folder_name in dirs:
            folder_path = os.path.join(root, folder_name)
            create_document_from_folder(folder_path, cover_path)

def remove_generated_files(base_path):
    for root, dirs, files in os.walk(base_path):
        for file in files:
            if file.endswith('_Solutions.docx'):
                file_path = os.path.join(root, file)
                os.remove(file_path)
                print(f'Removed file: {file_path}')

def select_directory():
    folder_path = filedialog.askdirectory(title="Select Base Directory")
    if folder_path:
        directory_var.set(folder_path)

def select_file():
    file_path = filedialog.askopenfilename(title="Select Cover Page File", filetypes=[("Word Documents", "*.docx")])
    if file_path:
        file_var.set(file_path)

def start_process():
    base_directory = directory_var.get()
    cover_file = file_var.get()
    
    if not base_directory or not cover_file:
        messagebox.showwarning("Input Error", "Please select both base directory and cover page file.")
        return
    
    process_all_folders(base_directory, cover_file)
    messagebox.showinfo("Success", "Documents have been created successfully.")

def revert_process():
    base_directory = directory_var.get()
    
    if not base_directory:
        messagebox.showwarning("Input Error", "Please select the base directory.")
        return
    
    remove_generated_files(base_directory)
    messagebox.showinfo("Revert Process", "All generated files have been removed.")

# Set up the Tkinter window
root = tk.Tk()
root.title("Document Generator")

tk.Label(root, text="Base Directory:").grid(row=0, column=0, padx=10, pady=10, sticky='e')
tk.Label(root, text="Cover Page File:").grid(row=1, column=0, padx=10, pady=10, sticky='e')

directory_var = tk.StringVar()
file_var = tk.StringVar()

tk.Entry(root, textvariable=directory_var, width=50).grid(row=0, column=1, padx=10, pady=10)
tk.Entry(root, textvariable=file_var, width=50).grid(row=1, column=1, padx=10, pady=10)

tk.Button(root, text="Browse...", command=select_directory).grid(row=0, column=2, padx=10, pady=10)
tk.Button(root, text="Browse...", command=select_file).grid(row=1, column=2, padx=10, pady=10)

tk.Button(root, text="Start Process", command=start_process).grid(row=2, column=1, padx=10, pady=20)
tk.Button(root, text="Revert Process", command=revert_process).grid(row=2, column=2, padx=10, pady=20)

root.mainloop()
