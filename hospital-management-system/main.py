from tkinter import *
from tkinter import messagebox,OptionMenu
from PIL import Image,ImageTk
import mysql
import mysql.connector
import smtplib as s
import random
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
def establish_connection():
    try:
        connection = mysql.connector.connect(host="localhost", user="root", password="", database="hospitalmanagementsystem")
        if connection.is_connected():
            return connection
    except mysql.connector.Error as err:
        messagebox.showerror("Error", f"Failed to establish connection: {err}")
def  addtotable():
    global Name, Age, Gender, Address, Email, Insuranceinfo, Medhistory, Status 
    conn = establish_connection()
    if conn:
        try:
            mycursor = conn.cursor()
            sql = f"SELECT  `DoctorID`  FROM `Doctors` WHERE `Name`='{doctorname.get()}'"
            mycursor.execute(sql)
            data=mycursor.fetchone()
            doctorid=data[0]
            conn.commit()
            sql = f"SELECT  *  FROM `appointments`"
            mycursor.execute(sql)
            data1=mycursor.fetchall()
            conn.commit()
            name1=Name.replace(' ',"_")
            id=f"SMPK-{name1.upper()}-{len(data1)+1}"
            characters = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789!@#$%^&*()_+,<.>/?;:"
            key = ''.join(random.sample(characters, 10))
            sql = f"INSERT INTO `appointments`(`patientid`,`password`, `name`, `age`, `address`,`gender`, `email`, `insuranceinfo`, `medhistory`, `doctorid`, `appointedday`, `appointedtime`,`status`) VALUES ('{id}','{key}','{Name}','{Age}','{Address}','{Gender}','{Email}','{Insuranceinfo}','{Medhistory}','{doctorid}','{day_var.get()}','{time_var.get()}',\"pending\")"
            mycursor.execute(sql)
            conn.commit()
            messagebox.showinfo("info","appointment is placed")
            frame1.destroy()
            main()
        except mysql.connector.Error as err:
            messagebox.showerror("Error", f"Error fetching doctor data: {err}")
        finally:
            conn.close()
def validate_otp():
    OTP=int(userotp.get())
    if (OTP==actualopt):
        messagebox.showinfo("info","mail is successfully verified")
        addtotable()
    else:
        messagebox.showerror("error","otp doesnot match")
def validate():
    global name, email, server,actualopt
    try:
        server = s.SMTP('smtp.gmail.com', 587)
        server.starttls()
        actualopt = random.randrange(100000, 999999)
        try:
            server.login('your mail', 'your password')
        except Exception as e:
            messagebox.showerror("error", "Failed to login")    
        if flag == 1:
            subjecttext = 'Approval for appointment'
            body = f'''
                Dear {patientdata[1]},
                it has found that patient with details:
                PatientID : {patientdata[0]}
                name : {patientdata[1]}
                Age : {patientdata[2]}
                Gender : {patientdata[3]}
                Address : {patientdata[4]}
                Insuranceinfo : {patientdata[6]}
                Medhistory : {patientdata[7]}
                
                has been appointed on this scheduled day and time: 

                day : {patientdata[-3]}
                time : {patientdata[-2]}
                regards
                {patientdata[-4]}
                SMPK ADMIN OFFICE
                '''
            message = f"Subject: {subjecttext}\n\n{body}"
            try:
                server.sendmail('your mail',patientdetails[5], message)
                messagebox.showinfo("info", "Successfully mail was sent")
                gofordoctor()
            except Exception as e:
                server.quit()
        elif flag == 2:
            subjecttext = 'prescription details'
            attachment_path ="medication_info.pdf"  
            body = f'''
                Dear {patientdetails[0]},
                YOUR APPOINTMENT HAS BEEN SUCCESSFULLY DONE. YOUR DETAILS AND PRESCRIPTION IS ATTACHED BELOW:
                Patient ID : {patientdetails[6]}
                Name : {patientdetails[0]}
                Age : {patientdetails[1]}
                Gender : {patientdetails[2]}
                Email : {patientdetails[3]}
                Doctor ID : {patientdetails[4]}
                Address : {patientdetails[5]}
                
                PLEASE FIND YOUR ATTACHMENT HERE
                regards
                SMPK HOSPITAL ADMIN OFFICE
                ANYWHERE IN THE UNIVERSE
                '''
            msg = MIMEMultipart()
            msg['From'] = 'your mail'
            msg['To'] = patientdetails[3]
            msg['Subject'] = subjecttext
            msg.attach(MIMEText(body, 'plain'))
            with open(attachment_path, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename= {os.path.basename(attachment_path)}')
            msg.attach(part)
            text = msg.as_string()
            try:
                server.sendmail('your mail', patientdetails[3], text)
                messagebox.showinfo("info", "Successfully mail was sent")
            except Exception as e:
                messagebox.showerror("error", f"Error in sending mail due to {e}")
                server.quit()
        elif(flag!=2 and flag!=1):
            subjecttext = 'Validation of Email'
            body = f'''
                Dear {name.get()},
                it is found that you are trying to take appointment for a doctor and first you have to verify your mail content. Your OTP is {actualopt}
                please enter this for proceeding further details
                Thank you for your cooperation.
                Sincerely,
                SMPK Hospital 
                '''
            message = f"Subject: {subjecttext}\n\n{body}"
            try:
                server.sendmail('your mail', email.get(), message)
                messagebox.showinfo("info", "Successfully mail was sent")
                gofordoctor()
            except Exception as e:
                messagebox.showerror("error", f"Error in sending mail due to {e}")
                server.quit()
    except Exception as e:
        messagebox.showerror("error", e)            
def gofordoctor():
    global name, email, age, gender, address, email, insuranceinfo, medhistory, status, type,userotp
    global Name, Age, Gender, Address, Email, Insuranceinfo, Medhistory, Status
    Name = name.get()
    Email = email.get()
    Address = address.get()
    Medhistory = medhistory.get("1.0", "end-1c")
    Status = status.get()
    Insuranceinfo = insuranceinfo.get()
    Gender = gender.get()
    Age = age.get()
    global frame1, b
    frame1.destroy()
    frame1 = Frame(root, bg=b)
    frame1.pack()
    img = Image.open(r"book-your-appointment-on-calendar-260nw-2305316035.webp").resize((800, 300))
    img = ImageTk.PhotoImage(img)
    label = Label(frame1, image=img)
    label.image = img
    label.grid(row=0, column=0, columnspan=5, padx=10, pady=10)
    conn = establish_connection()
    if conn:
        try:
            mycursor = conn.cursor()
            sql = f"SELECT  `DoctorID`, `Name`,  `Working_days`, `Schedule_time`  FROM `Doctors` WHERE `Specialization`='{type.get()}'"
            mycursor.execute(sql)
            data = mycursor.fetchall()
            conn.commit()
            doctornames = []
            schedule = []
            timing = []
            for i in data:
                doctornames.append(i[1])
                schedule.append(i[-2])
                timing.append(i[-1])
            global doctorname
            Label(frame1, text="Select the doctor", font="stencil 20 bold", bg=b, fg="white").grid(row=1, column=0, columnspan=10, padx=10, pady=10)
            doctorname = StringVar()
            doctor_menu = OptionMenu(frame1, doctorname, *doctornames, command=update_schedule_and_timing)
            doctor_menu.grid(row=2, column=1, padx=10, pady=10)
            doctorname.set("Choose")
            Label(frame1, text="Specialist:", bg=b, fg="white", font="stencil 24 bold").grid(row=2, column=0, padx=10, pady=10)
        except mysql.connector.Error as err:
            messagebox.showerror("Error", f"Error fetching doctor data: {err}")
        finally:
            conn.close()
    Label(frame1, text="Enter OTP:", font="stencil 20 bold", bg=b, fg="white").grid(row=3, column=0)
    userotp = Entry(frame1, width=30, fg=b)
    userotp.grid(row=3, column=1)
def update_schedule_and_timing(selected_doctor):
    conn = establish_connection()
    if conn:
        try:
            mycursor = conn.cursor()
            sql = "SELECT `Working_days`, `Schedule_time` FROM `Doctors` WHERE `Name`=%s"
            mycursor.execute(sql, (selected_doctor,))
            data = mycursor.fetchone()
            conn.commit()
            if data:
                working_days = data[0]
                schedule_time = data[1]
                working_days_list = working_days.split(',')
                schedule_time_list = schedule_time.split(',')
                update_day_dropdown(working_days_list)
                update_time_dropdown(schedule_time_list)
        except mysql.connector.Error as err:
            messagebox.showerror("Error", f"Error fetching doctor's schedule and timing: {err}")
        finally:
            conn.close()
def update_day_dropdown(working_days_list):
    global frame1
    Label(frame1, text="Select Day:", font="stencil 20 bold", bg=b, fg="white").grid(row=4, column=0)
    global day_var
    day_var = StringVar()
    day_dropdown = OptionMenu(frame1, day_var, *working_days_list)
    day_var.set("choose day")
    day_dropdown.grid(row=4, column=1, padx=10, pady=10)
def update_time_dropdown(schedule_time_list):
    global frame1
    Label(frame1, text="Select Time:", font="stencil 20 bold", bg=b, fg="white").grid(row=5, column=0)
    all_time_slots = []
    for time_interval in schedule_time_list:
        start_time_str, end_time_str = time_interval.split('-')
        start_time = datetime.strptime(start_time_str.strip(), '%H:%M')
        end_time = datetime.strptime(end_time_str.strip(), '%H:%M')
        current_time = start_time
        while current_time <= end_time:
            all_time_slots.append(current_time.strftime('%I:%M %p'))
            current_time += timedelta(minutes=30)
    global time_var
    time_var = StringVar()
    time_dropdown = OptionMenu(frame1, time_var, *all_time_slots)
    time_var.set("select time")
    time_dropdown.grid(row=5, column=1, padx=10, pady=10)
    img=Image.open("book-your-appointment-on-calendar-260nw-2305316035.webp").resize((200,50))
    img=ImageTk.PhotoImage(img)
    label=Button(frame1,image=img,command=validate_otp)
    label.image=img
    label.grid(row=6,column=0,columnspan=10)
def register():
    global frame1,b,name,age,gender,address,email,insuranceinfo,medhistory,status,type,flag
    flag=0
    frame1.destroy()
    frame1=Frame(root,bg=b)
    frame1.pack()
    img=Image.open(r"book-your-appointment-on-calendar-260nw-2305316035.webp").resize((800,300))
    img=ImageTk.PhotoImage(img)
    label=Label(frame1,image=img)
    label.image=img
    label.grid(row=0,column=0,columnspan=5)
    Label(frame1,text="enter patient details",font="stencil 30 bold",justify="center",bg=b,fg="#ffccbb").grid(row=1,column=0,columnspan=5)
    Label(frame1,text="Name : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=2,column=0,padx=10,pady=10)
    name=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
    name.grid(row=2,column=1)
    Label(frame1,text="Age : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=2,column=2,padx=10,pady=10)
    age=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
    age.grid(row=2,column=3)
    Label(frame1,text="Gender : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=3,column=0,padx=10,pady=10)
    gender=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
    gender.grid(row=3,column=1)
    Label(frame1,text="Address : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=3,column=2,padx=10,pady=10)
    address=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
    address.grid(row=3,column=3)
    Label(frame1,text="Email : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=4,column=0,padx=10,pady=10)
    email=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
    email.grid(row=4,column=1)
    Label(frame1,text="Insurance info : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=4,column=2,padx=10,pady=10)
    insuranceinfo=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
    insuranceinfo.grid(row=4,column=3)
    Label(frame1,text="medical history : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=5,column=0,padx=10,pady=10)
    medhistory=Text(frame1,fg=b,font="helvetica 12 italic",width=70,height=5)
    medhistory.grid(row=5,column=1,columnspan=3,padx=10,pady=10)
    values=['General Practitioner','Pediatrician','Cardiologist','Dermatologist','Ophthalmologist','Psychiatrist','Neurologist','Orthopedic Surgeon','Gynecologist','Urologist','Pulmonologist','Endocrinologist','Gastroenterologist','Rheumatologist','Hematologist','Infectious Disease Specialist','Nephrologist','Allergist','Plastic Surgeon']
    type=StringVar()
    type.set("choose")
    OptionMenu(frame1,type,*values).grid(row=6,column=1,padx=10,pady=10)
    Label(frame1,text="Specialist :",bg=b,fg="white",font="stencil 24 bold").grid(row=6,column=0,padx=10,pady=10)
    Label(frame1,text="Status : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=6,column=2,padx=10,pady=10)
    status=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
    status.insert(0,"Pending")
    status.grid(row=6,column=3)
    status.config(state="readonly")
    img1=Image.open(r"book-your-appointment-on-calendar-260nw-2305316035.webp").resize((200,50))
    img1=ImageTk.PhotoImage(img1)
    btn1=Button(frame1,image=img1,command=validate)
    btn1.image=img1
    btn1.grid(row=7,column=0,columnspan=5)
def appoint():
    conn=establish_connection()
    try:
        mycursor= conn.cursor()
        sql=f"UPDATE `appointments` SET `status`='appointed' WHERE `patientid`='{patientdata[0]}';"
        mycursor.execute(sql)
        conn.commit()
        conn.close()
        global flag
        flag=1
        validate()
        messagebox.showinfo("info","appointment booked successfully")
        frame1.destroy()
        main()
    except Exception as e:
        messagebox.showerror("error","failed to  appoint ")
def details():
    conn=establish_connection()
    try:
        global frame1,patientdata
        sql=f"SELECT `patientid`, `name`, `age`,`gender`, `address`,`email`, `insuranceinfo`, `medhistory`, `doctorid`, `appointedday`, `appointedtime`, `status` FROM `appointments` WHERE `name` = '{patient.get()}'"
        frame1.destroy()
        frame1=Frame()
        frame1.pack()
        frame1.configure(bg=b)
        mycursor=conn.cursor()
        mycursor.execute(sql)
        patientdata=mycursor.fetchone()
        conn.commit()
        conn.close()
        img=Image.open(r"book-your-appointment-on-calendar-260nw-2305316035.webp").resize((800,300))
        img=ImageTk.PhotoImage(img)
        label=Label(frame1,image=img)
        label.image=img
        label.grid(row=0,column=0,columnspan=5)
        Label(frame1,text=patientdata[0],font="stencil 30 bold",justify="center",bg=b,fg="#ffccbb").grid(row=1,column=0,columnspan=5)
        Label(frame1,text="Name : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=2,column=0,padx=10,pady=10)
        name=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        name.grid(row=2,column=1)
        name.insert(0,patientdata[1])
        name.config(state="readonly")
        Label(frame1,text="Age : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=2,column=2,padx=10,pady=10)
        age=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        age.grid(row=2,column=3)
        age.insert(0,patientdata[2])
        age.config(state="readonly")
        Label(frame1,text="Gender : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=3,column=0,padx=10,pady=10)
        gender=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        gender.grid(row=3,column=1)
        gender.insert(0,patientdata[3])
        gender.config(state="readonly")
        Label(frame1,text="Address : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=3,column=2,padx=10,pady=10)
        address=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        address.grid(row=3,column=3)
        address.insert(0,patientdata[4])
        address.config(state="readonly")
        Label(frame1,text="Email : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=4,column=0,padx=10,pady=10)
        email=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        email.grid(row=4,column=1)
        email.insert(0,patientdata[5])
        email.config(state="readonly")
        Label(frame1,text="Insurance info : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=4,column=2,padx=10,pady=10)
        insuranceinfo=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        insuranceinfo.grid(row=4,column=3)
        insuranceinfo.insert(0,patientdata[6])
        insuranceinfo.config(state="readonly")
        Label(frame1,text="medical history : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=5,column=0,padx=10,pady=10)
        medhistory=Text(frame1,fg=b,font="helvetica 12 italic",width=80,height=3)
        medhistory.grid(row=5,column=1,columnspan=3,padx=10,pady=10)
        medhistory.insert(END,patientdata[7])
        medhistory.config(state="disabled") 
        Label(frame1,text="doctorid :",bg=b,fg="white",font="stencil 24 bold").grid(row=6,column=0,padx=10,pady=10)
        type=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        type.grid(row=6,column=1)
        type.insert(0,patientdata[-4])
        type.config(state="readonly")
        Label(frame1,text="Status : ",font="stencil 20 bold",justify="left",bg=b,fg="white").grid(row=6,column=2,padx=10,pady=10)
        status=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        status.grid(row=6,column=3)
        status.insert(0,patientdata[-1])
        Label(frame1,text=" Appointment day :",bg=b,fg="white",font="stencil 24 bold").grid(row=7,column=0,padx=10,pady=10)
        type=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        type.grid(row=7,column=1)
        type.insert(0,patientdata[-3])
        type.config(state="readonly")
        Label(frame1,text="Appointed time :",bg=b,fg="white",font="stencil 24 bold").grid(row=7,column=2,padx=10,pady=10)
        type=Entry(frame1,fg=b,font="helvetica 12 italic",width=30)
        type.grid(row=7,column=3)
        type.insert(0,patientdata[-2])
        type.config(state="readonly")
        img1=Image.open(r"book-your-appointment-on-calendar-260nw-2305316035.webp").resize((200,50))
        img1=ImageTk.PhotoImage(img1)
        btn1=Button(frame1,image=img1,command=appoint)
        btn1.image=img1
        btn1.grid(row=8,column=0,columnspan=5)
    except Exception as e:
        pass
def goforlogin():
    global frame1
    conn=establish_connection()
    try:
        mycursor=conn.cursor()
        sql=f"select `Name` from `Doctors` where `doctorId`=\"{doctorid.get()}\" and `password`=\"{key.get()}\""
        mycursor.execute(sql)
        data=mycursor.fetchone()
        conn.commit()
        sql=f"SELECT `name`  FROM `appointments` WHERE `doctorid`=\"{doctorid.get()}\" AND `status`='pending'"
        mycursor.execute(sql)
        data=mycursor.fetchall()
        conn.commit()
        patients=[]
        for i in data:
            patients.append(i[0])
        if len(data) ==0 :
            messagebox.showinfo("info","no more appointments")
        else:
            try:
                global b
                frame1.destroy()
                frame1=Frame(root,bg=b)
                frame1.pack()
                Label(frame1, text="Select the patients", font="stencil 20 bold", bg=b, fg="white").grid(row=0, column=0, columnspan=10, padx=10, pady=10)
                Label(frame1, text="patient name :", bg=b, fg="light pink", font="stencil 24 bold").grid(row=1, column=0, padx=10, pady=10)
                global patient
                patient = StringVar()
                patient_menu = OptionMenu(frame1, patient, *patients, command=None)
                patient_menu.grid(row=1, column=1, padx=10, pady=10)      
                img1=Image.open(r"login_nl_cover.jpg").resize((200,80))
                img1=ImageTk.PhotoImage(img1)
                btn1=Button(frame1,image=img1,command=details)
                btn1.image=img1
                btn1.grid(row=4,columnspan=10)
            except Exception as e:
                pass
    except Exception as e:
        messagebox.showerror("error","invalid login credentials")
def doctorlogin():
    global frame1,b,frame2,doctorid,key
    frame1.destroy()
    frame1=Frame(root)
    frame1.pack()
    frame1.configure(bg=b)
    img=Image.open(r"_3b3a3987-4b41-4347-a82e-889fe8461441.jpg").resize((500,400))
    img=ImageTk.PhotoImage(img)
    label=Label(frame1,bg=b,image=img)
    label.iamge=img
    frame1.configure(bg=b)
    label.grid(row=0,column=0,columnspan=10)
    Label(frame1,text="doctor login",font="stencil 20 bold",bg=b,fg="white",justify="center").grid(row=1,columnspan=10)
    Label(frame1,text="doctor UID",font="jokerman 15 italic",bg=b,fg="#ccff23",justify="center").grid(row=2,column=0)
    doctorid=Entry(frame1,width=30,fg=b,font="helvetica 12 italic")
    doctorid.grid(row=2,column=1)
    Label(frame1,text="Password : ",font="jokerman 15 italic",bg=b,fg="#ccff23",justify="center").grid(row=3,column=0)
    key=Entry(frame1,width=30,fg=b,font="helvetica 12 italic",show="*")
    key.grid(row=3,column=1)
    img1=Image.open(r"login_nl_cover.jpg").resize((200,80))
    img1=ImageTk.PhotoImage(img1)
    btn1=Button(frame1,image=img1,command=goforlogin)
    btn1.image=img1
    btn1.grid(row=4,columnspan=10)
def update_datetime():
    date = datetime.now()
    currentdate.config(text="Current Date: " + date.strftime("%Y-%m-%d"))
    currenttime.config(text="Current Time: " + date.strftime("%I:%M:%S %p")) 
    root.after(1000, update_datetime)
def createattachment():
    def add_medication_info(doc, medicines, doses, durations, frequencies, routes, instructions, notes):
        medicines = medicines
        doses = doses
        durations = durations
        frequencies =frequencies
        routes = routes
        instructions = instructions
        notes = notes
        heading = doc.add_heading('SMPK HOSPITAL', level=1)
        run = heading.runs[0]
        run.font.name = 'Stencil'
        run.font.size = Pt(30)
        run.bold = True
        doc.add_picture('hidden_leaf_village__hospital__3_by_iennidesign_d8bylt1-pre.jpg', width=Inches(7))
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(f'Patient Details', level=2)
        doc.add_paragraph(f'Patient ID : {patientdetails[6]}')
        doc.add_paragraph(f'Name : {patientdetails[0]}')
        doc.add_paragraph(f'Age : {patientdetails[1]}')
        doc.add_paragraph(f'Gender : {patientdetails[2]}')
        doc.add_paragraph(f'Email : {patientdetails[3]}')
        doc.add_paragraph(f'Doctor ID : {patientdetails[4]}')
        doc.add_paragraph(f'Address : {patientdetails[5]}')
        doc.add_heading('Medication Information', level=1)
        for i in range(len(medicines)-1):
            doc.add_heading(f'Medication {i+1}: {medicines[i]}', level=2)
            doc.add_paragraph(f'Dose: {doses[i]}ml ')
            doc.add_paragraph(f'Duration: {durations[i]} hour')
            doc.add_paragraph(f'Frequency: {frequencies[i]} times per day')
            doc.add_paragraph(f'Route: {routes[i]}')
            doc.add_paragraph(f'Instructions: {instructions[i]}')
            doc.add_paragraph(f'Notes : {notes[i]}')
            if i < len(medicines) - 1:
                pass
    doc = Document()
    add_medication_info(doc, medicines, doses, durations, frequencies, routes, instructions, notes)
    doc.add_heading('Discription', level=1)
    doc.add_paragraph(f'{" ".join(notes[-1])}')
    doc.save('medication_info.docx')
    convert("medication_info.docx")
def sendprescription():
    global flag
    flag=2
    validate()
def prescribed():
    conn=establish_connection()
    pid_str = pid
    docid_str = docid
    medicines_str = ', '.join(map(str, medicines))
    doses_str = ', '.join(map(str, doses))
    durations_str = ', '.join(map(str, durations))
    frequencies_str = ', '.join(map(str, frequencies))
    routes_str = ', '.join(map(str, routes))
    instructions_str = ', '.join(map(str, instructions))
    notes_str = ', '.join(map(str, notes))    
    mycursor= conn.cursor()
    sql=f"UPDATE `appointments` SET `status`='prescribed' WHERE `patientid`='{pid_str}';"
    mycursor.execute(sql)
    conn.commit()
    sql = f"INSERT INTO `prescription` (`patient_id`, `doctor_key`, `medicines`, `doses`, `durations`, `frequencies`, `routes`, `instructions`, `notes`) VALUES ('{pid_str}', '{docid_str.replace('\'', '\'\'')}', '{medicines_str.replace('\'', '\'\'')}', '{doses_str.replace('\'', '\'\'')}', '{durations_str.replace('\'', '\'\'')}', '{frequencies_str.replace('\'', '\'\'')}', '{routes_str.replace('\'', '\'\'')}', '{instructions_str.replace('\'', '\'\'')}', '{notes_str.replace('\'', '\'\'')}')"
    try:
        mycursor=conn.cursor()
        mycursor.execute(sql)
        conn.commit()
        conn.close()
        messagebox.showinfo("info","prescribed successsfully")
        try:
            createattachment()
        except Exception as e:
            messagebox.showerror("error",e)
        messagebox.showinfo("info","attachment created")
        sendprescription()
        frame1.destroy()
        main()
    except Exception as e:
        messagebox.showerror("error","failed to prescribe\ntry again later")
def doctorprescribe():
    global medicines,doses,frequencies,routes,durations,instructions,frame1,currentdate,currenttime,docid,pid,notes,patientdetails
    pid=prescribe_id.get()
    prescribe_key.get()
    conn=establish_connection()
    try:
        mycursor=conn.cursor()
        sql1=f"SELECT  `name`, `age`, `gender`, `email`,  `doctorid`, `address` ,`patientid` FROM `appointments` WHERE `doctorid`='{doctorid.get()}' and `doctorpass`='{prescribe_key.get()}'and `patientid`='{prescribe_id.get()}' "
        mycursor.execute(sql1)
        patientdetails=mycursor.fetchone()
        conn.commit()
        conn.close()
        frame1.destroy()
        frame1=Frame(root,bg=b)
        frame1.pack()
        Label(frame1,text="patient details",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=0,column=0,columnspan=10,padx=10,pady=10)
        Label(frame1,text="Name : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=0,padx=10,pady=10)
        Label(frame1,text=patientdetails[0],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=1,padx=10,pady=10)
        Label(frame1,text="Age  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=2,padx=10,pady=10)
        Label(frame1,text=patientdetails[1],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=3,padx=10,pady=10)
        Label(frame1,text="Sex  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=4,padx=10,pady=10)
        Label(frame1,text=patientdetails[2],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=5,padx=10,pady=10)
        docid=patientdetails[4]
        pid=patientdetails[-1]
        Label(frame1,text="Email : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=0,padx=10,pady=10)
        Label(frame1,text=patientdetails[3],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=1,padx=10,pady=10)
        Label(frame1,text="doctorid  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=2,padx=10,pady=10)
        Label(frame1,text=patientdetails[4],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=3,padx=10,pady=10)
        Label(frame1,text="address  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=4,padx=10,pady=10)
        Label(frame1,text=patientdetails[5],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=5,padx=10,pady=10)
    except EXCEPTION as e:
        messagebox.showerror("error","failed to establish connectioin")
    #date
    currentdate = Label(frame1,bg=b,font="helvetica 30 italic",fg="light pink",justify="left")
    currentdate.grid(row=3, column=0, padx=20, pady=(30, 5), sticky=W,columnspan=3)
    currenttime = Label(frame1,bg=b,font="helvetica 30 italic",fg="light pink",justify="right")
    currenttime.grid(row=3, column=3, padx=20, pady=(30, 5), sticky=W,columnspan=3)
    update_datetime()
    ''''Prescription Details:
    Name of the medication: The generic or brand name of the prescribed medication.
    Dosage: The amount of medication to be taken (e.g., mg, ml).
    Frequency: How often the medication should be taken (e.g., once daily, twice daily).
    Route: The method by which the medication should be administered (e.g., orally, intravenously).
    Duration: The length of time for which the medication should be taken.
    Instructions: Any specific instructions for taking the medication (e.g., with food, before bedtime).'''
    Label(frame1,text="medications",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=4,column=0,columnspan=10,padx=10,pady=10)
    Label(frame1,text="medicines : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=5,column=0,padx=10,pady=10)
   
    
    Label(frame1,text="medicines : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=5,column=0,padx=10,pady=10)
    medicines=[]
    medicine=Entry(frame1,width=30,fg=b)
    medicine.grid(row=5,column=1,padx=10,pady=10)
    
    Label(frame1,text="doses : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=5,column=2,padx=10,pady=10)
    doses=[]
    dose=Entry(frame1,width=30,fg=b)
    dose.grid(row=5,column=3,padx=10,pady=10)
    
    Label(frame1,text="frequency : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=5,column=4,padx=10,pady=10)
    frequencies=[]
    frequecy=Entry(frame1,width=30,fg=b)
    frequecy.grid(row=5,column=5,padx=10,pady=10)
    
    Label(frame1,text="Route : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=6,column=0,padx=10,pady=10)
    routes=[]
    route=Entry(frame1,width=30,fg=b)
    route.grid(row=6,column=1,padx=10,pady=10)
    
    Label(frame1,text="duration : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=6,column=2,padx=10,pady=10)
    durations=[]
    duration=Entry(frame1,width=30,fg=b)
    duration.grid(row=6,column=3,padx=10,pady=10)
    
    Label(frame1,text="instructions : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=6,column=4,padx=10,pady=10)
    instructions=[]
    instruction=Entry(frame1,width=30,fg=b)
    instruction.grid(row=6,column=5,padx=10,pady=10)
    
    Label(frame1,text="notes :",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=7,column=0,padx=10,pady=10)
    notes=[]
    note=Text(frame1,fg=b,font="helvetica 12 italic",width=70,height=5)
    note.grid(row=7,column=1,columnspan=4,padx=10,pady=10)
    def add():
        medicines.append(medicine.get())
        doses.append(dose.get())
        durations.append(duration.get())
        frequencies.append(frequecy.get())
        routes.append(route.get())
        instructions.append(instruction.get())
        notes.append(note.get("1.0",END))
        medicine.delete(0,END)
        dose.delete(0,END)
        duration.delete(0,END)
        frequecy.delete(0,END)
        route.delete(0,END)
        instruction.delete(0,END)
        note.delete("1.0",END)
    img1=Image.open(r"seo.png").resize((200,80))
    img1=ImageTk.PhotoImage(img1)
    btn1=Button(frame1,image=img1,command=add)
    btn1.image=img1
    btn1.grid(row=8,column=1)

    img1=Image.open(r"_b06199d5-b8c1-4d57-80d8-f17b93c23fa7.jpg").resize((200,80))
    img1=ImageTk.PhotoImage(img1)
    btn1=Button(frame1,image=img1,command=prescribed)
    btn1.image=img1
    btn1.grid(row=8,column=4)

def patientprescribe():
    global medicines,doses,frequencies,routes,durations,instructions,frame1,currentdate,currenttime,docid,pid,notes
    pid=prescribe_id.get()
    prescribe_key.get()
    #patient details
    conn=establish_connection()
    try:
        mycursor=conn.cursor()
        sql1=f"SELECT  `name`, `age`, `gender`, `email`,  `doctorid`, `address` ,`patientid` FROM `appointments` WHERE `doctorid`='{doctorid.get()}' and `password`='{prescribe_key.get()}'and `patientid`='{prescribe_id.get()}' "
        mycursor.execute(sql1)
        patientdetails=mycursor.fetchone()
        conn.commit()
        conn.close()
        frame1.destroy()
        frame1=Frame(root,bg=b)
        frame1.pack()
        Label(frame1,text="patient details",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=0,column=0,columnspan=10,padx=10,pady=10)
        Label(frame1,text="Name : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=0,padx=10,pady=10)
        Label(frame1,text=patientdetails[0],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=1,padx=10,pady=10)
        Label(frame1,text="Age  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=2,padx=10,pady=10)
        Label(frame1,text=patientdetails[1],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=3,padx=10,pady=10)
        Label(frame1,text="Sex  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=4,padx=10,pady=10)
        Label(frame1,text=patientdetails[2],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=5,padx=10,pady=10)
        docid=patientdetails[4]
        Label(frame1,text="Email : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=0,padx=10,pady=10)
        Label(frame1,text=patientdetails[3],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=1,padx=10,pady=10)
        Label(frame1,text="doctorid  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=2,padx=10,pady=10)
        Label(frame1,text=patientdetails[4],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=3,padx=10,pady=10)
        Label(frame1,text="address  : ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=4,padx=10,pady=10)
        Label(frame1,text=patientdetails[5],bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=2,column=5,padx=10,pady=10)
        currentdate = Label(frame1,bg=b,font="helvetica 30 italic",fg="light pink",justify="left")
        currentdate.grid(row=3, column=0, padx=20, pady=(30, 5), sticky=W,columnspan=3)
        currenttime = Label(frame1,bg=b,font="helvetica 30 italic",fg="light pink",justify="right")
        currenttime.grid(row=3, column=3, padx=20, pady=(30, 5), sticky=W,columnspan=3)
        update_datetime()
        conn=establish_connection()
        mycursor = conn.cursor()
        sql = f"SELECT `medicines`, `doses`, `durations`, `frequencies`, `routes`, `instructions`, `notes` FROM `prescription` WHERE `patient_id`='{pid}'"
        mycursor.execute(sql)
        data = mycursor.fetchone()
        if data:
            medicines, doses, durations, frequencies, routes, instructions, notes = data
        conn.commit()
        conn.close()
        medicines=medicines.split(",")
        doses=doses.split(",")
        durations=durations.split(",")
        frequencies=frequencies.split(",")
        routes=routes.split(",")
        instructions=instructions.split(",")
        Label(frame1,text="prescription",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=4,column=0,columnspan=10,padx=10,pady=10)
        Label(frame1,text="s.no.",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=5,column=0,padx=10,pady=10)

        Label(frame1,text="Medicines",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=5,column=1,padx=10,pady=10)
        Label(frame1,text="Doses",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=5,column=2,padx=10,pady=10)
        Label(frame1,text="durations",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=5,column=3,padx=10,pady=10)
        Label(frame1,text="frequencies",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=5,column=4,padx=10,pady=10)
        Label(frame1,text="routes",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=5,column=5,padx=10,pady=10)
        Label(frame1,text="instructions",bg=b,fg="#ccff23",font="stencil 20 bold",justify="center").grid(row=5,column=6,padx=10,pady=10)
        row=6
        for i in range(len(data[0].split(","))):
            Label(frame1,text=i+1,bg=b,fg="#ccff23",font="helvetica 12 italic",justify="center").grid(row=row+i,column=0,padx=10,pady=10)
            Label(frame1,text=medicines[i],bg=b,fg="#ccff23",font="helvetica 12 italic",justify="center").grid(row=row+i,column=1,padx=10,pady=10)
            Label(frame1,text=doses[i],bg=b,fg="#ccff23",font="helvetica 12 italic",justify="center").grid(row=row+i,column=2,padx=10,pady=10)
            Label(frame1,text=durations[i],bg=b,fg="#ccff23",font="helvetica 12 italic",justify="center").grid(row=row+i,column=3,padx=10,pady=10)
            Label(frame1,text=frequencies[i],bg=b,fg="#ccff23",font="helvetica 12 italic",justify="center").grid(row=row+i,column=4,padx=10,pady=10)
            Label(frame1,text=routes[i],bg=b,fg="#ccff23",font="helvetica 12 italic",justify="center").grid(row=row+i,column=5,padx=10,pady=10)
            Label(frame1,text=instructions[i],bg=b,fg="#ccff23",font="helvetica 12 italic",justify="center").grid(row=row+i,column=6,padx=10,pady=10)
            r=row+i
        Label(frame1,text='notes',bg=b,fg="#ccff23",font="stencil 20 italic",justify="center").grid(row=r+2,column=0,padx=10,pady=10)
        txt=Text(frame1,font="helvetica 12 italic",width=70,height=5,fg=b)
        txt.grid(row=r+2,column=1,columnspan=9)
        txt.insert('1.0',notes.replace(","," "))
        txt.config(state="disabled")
        def home():
            frame1.destroy()
            main()
        img1=Image.open(r"1024px-Go-home.svg.png").resize((200,80))
        img1=ImageTk.PhotoImage(img1)
        btn1=Button(frame1,image=img1,command=home)
        btn1.image=img1
        btn1.grid(row=r+3,columnspan=10)
    except Exception as e:
        messagebox.showerror("error","failed to establish connectioin")
    

def prescribe():
    if (canditate.get() == "patients"):
        patientprescribe()
    elif (canditate.get() == "doctors"):
        doctorprescribe()
def choose():
    global frame1,b,doctorid,canditate,prescribe_id,prescribe_key
    frame1.destroy()
    frame1=Frame(root,bg=b)
    frame1.pack()
    frame1.configure(bg=b)
    img=Image.open(r"_3b3a3987-4b41-4347-a82e-889fe8461441.jpg").resize((500,400))
    img=ImageTk.PhotoImage(img)
    label=Label(frame1,bg=b,image=img)
    label.image=img
    frame1.configure(bg=b)
    label.grid(row=0,column=0,columnspan=10,padx=10,pady=10)
    Label(frame1,text="   Role   ",bg=b,fg="#ccff23",font="jokerman 20 bold",justify="center").grid(row=1,column=0,padx=10,pady=10)
    values=["patients","doctors"]
    canditate=StringVar()
    option_menu=OptionMenu(frame1,canditate,*values)
    option_menu.grid(row=1,column=1,padx=10,pady=10)
    option_menu.config(width=30, bg="lightblue", fg="light pink",font="helvetica 15 italic")
    canditate.set("Choose your role")
    
    Label(frame1,text="Doctor ID : ",font="jokerman 15 italic",bg=b,fg="#ccff23",justify="center").grid(row=2,column=0)
    doctorid=Entry(frame1,width=30,fg=b,font="helvetica 12 italic")
    doctorid.grid(row=2,column=1)
    
    Label(frame1,text="Patient ID : ",font="jokerman 15 italic",bg=b,fg="#ccff23",justify="center").grid(row=3,column=0)
    prescribe_id=Entry(frame1,width=30,fg=b,font="helvetica 12 italic")
    prescribe_id.grid(row=3,column=1)
    Label(frame1,text="Password : ",font="jokerman 15 italic",bg=b,fg="#ccff23",justify="center").grid(row=4,column=0)
    prescribe_key=Entry(frame1,width=30,fg=b,font="helvetica 12 italic",show="*")
    prescribe_key.grid(row=4,column=1)
    #book appointment
    img1=Image.open(r"login_nl_cover.jpg").resize((200,80))
    img1=ImageTk.PhotoImage(img1)
    btn1=Button(frame1,image=img1,command=prescribe)
    btn1.image=img1
    btn1.grid(row=5,columnspan=10)
def main():
    try:
        os.remove("medication_info.pdf")
        os.remove("medication_info.docx")
    except :
        pass
    global frame1,b,frame2
    frame1=Frame(root)
    frame1.pack()
    frame1.configure(bg=b)
    frame2=Frame(root)
    frame2.pack()
    Label(frame1,text="WELCOME TO SMPK HOSPITAL",font="stencil 50 bold",bg=b,fg="silver").grid(row=0,column=0,columnspan=10)
    img=Image.open(r"hidden_leaf_village__hospital__3_by_iennidesign_d8bylt1-pre.jpg").resize((500,300))
    img=ImageTk.PhotoImage(img)
    label=Label(frame1,bg=b,image=img)
    label.iamge=img
    label.grid(row=1,column=0,columnspan=10)
    Label(frame1, text="""Welcome to SMPK Hospital! We understand that your health and well-being are of utmost importance. That's why we're dedicated to providing you with top-notch medical care and a comfortable environment during your  stay. Our hospital offers a wide range of services, from routine check ups to specialized treatments, all delivered by our skilled medical professionals. Whether you're here for a quick consultation or a longer stay, we have options to suit your needs. Our patient rooms are designed to ensure your comfort, with amenities such as complimentary Wi-Fi and attentive nursing staff.We also offer flexible appointment scheduling and personalized treatment plans to accommodate your preferences and health requirements.SMPK Hospital is committed to your well-being.Thank you for choosing us.We look forward to providing you with a positive and healing experience.""",font="Courier 15 italic", fg="#FBFFB1",justify='center',wraplength=1100,bg=b).grid(row=2, column=0, columnspan=10)
    Label(frame1,text="Our services",bg=b,fg="#222831",font="Roman 50 italic",justify='center').grid(row=3,columnspan=10,sticky=N)
    #book appointment
    img1=Image.open(r"book-your-appointment-on-calendar-260nw-2305316035.webp").resize((200,100))
    img1=ImageTk.PhotoImage(img1)
    btn1=Button(frame1,image=img1,command=register)
    btn1.image=img1
    btn1.grid(row=4,column=0)
    #grant appointment
    img2=Image.open(r"_22e05599-005e-41b9-9e6f-be7df352b50d.jpg").resize((200,100))
    img2=ImageTk.PhotoImage(img2)
    btn2=Button(frame1,image=img2,command=doctorlogin)
    btn2.image=img2
    btn2.grid(row=4,column=3)
    #prescribe
    img2=Image.open(r"AdobeStock_300785040.jpeg").resize((200,100))
    img2=ImageTk.PhotoImage(img2)
    btn2=Button(frame1,image=img2,command=choose)
    btn2.image=img2
    btn2.grid(row=4,column=6)
root=Tk()
root.title("Hospital Management system")
root.iconbitmap(r"hospital_health_clinic_urban_buildings_medical_icon_128575.ico")
root.geometry("+0+0")
b="#B15EFF"
main()
root.mainloop()